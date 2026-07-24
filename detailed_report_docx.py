"""
detailed_report_docx.py — 詳細輿情分析報告（Word 版）。

與 detailed_report_generator（純文字 .md）內容一致，但改用真正的 Word 標題與表格，
供正式交付。沿用 framework_generator 的 docx 排版元件（微軟正黑體、表頭底色、斑馬紋）。
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

import framework_generator as fw
from detailed_report_generator import (
    _parse_ai, _strip_md, _clean, _fmt_date, _growth_lines, _safe_float,
    AI_CAPTION_NOTE,
)
from keypo_loader import classify_article, match_keywords_to_articles

FONT = fw.FONT
C_INK, C_GRAY, C_RED = fw.C_INK, fw.C_GRAY, fw.C_RED


def _meta_line(doc, label, value):
    p = fw._para(doc, "", space_after=2)
    fw._font(p.add_run(f"{label}："), size=10, bold=True, color=C_GRAY)
    fw._font(p.add_run(str(value)), size=10, color=C_INK)


def _caption(doc, text):
    """AI 敘事摘要（灰字小號）+ 核對提醒，對照 .md 版的表格摘要句"""
    if not text:
        return
    fw._para(doc, str(text), size=10, color=C_GRAY, space_before=4, space_after=1)
    fw._para(doc, AI_CAPTION_NOTE, size=8.5, color=C_GRAY, space_after=2)


def generate_detailed_report_docx(
    brand, sentiment, keywords, articles, trend, kols, ai_report,
    start_date, end_date, out_path,
    growth_info=None, competitor_data=None, table_captions=None,
    source_dist=None, trend_narrative=None,
):
    table_captions = table_captions or {}
    disp_s, disp_e = _fmt_date(start_date), _fmt_date(end_date)
    period = f"{disp_s} – {disp_e}"
    days = len(trend["daily"]) or 1
    avg_vol = trend["total"] // days
    ai = _parse_ai(ai_report)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # ── 標題 ─────────────────────────────────────────
    t = fw._para(doc, f"{brand} 品牌網路輿情分析報告", size=22, bold=True,
                 color=C_INK, space_after=2)
    fw._red_rule(t, sz=20)
    fw._para(doc, "詳細分析報告", size=13, color=C_RED, space_after=8)
    _meta_line(doc, "觀測期間", f"{period}（共 {days} 天）")
    _meta_line(doc, "資料來源", "KEYPO 大數據關鍵引擎")
    _meta_line(doc, "產出日期", datetime.today().strftime("%Y/%m/%d"))
    _meta_line(doc, "分析工具", "InsightFlow AI 輿情分析系統")

    # ── §1 執行摘要 ──────────────────────────────────
    fw._section(doc, "§1　執行摘要")
    summary = _strip_md(ai.get("執行摘要", "")).strip()
    if summary:
        fw._para(doc, summary)
    concl = [
        f"本週總聲量 {trend['total']:,} 筆，平均每日 {avg_vol:,} 筆。",
        f"聲量高峰：{trend['peak_date']}（{trend['peak_volume']:,} 筆）。",
        f"情緒結構：正面 {sentiment['positive_percent']}、中立 "
        f"{sentiment['neutral_percent']}、負面 {sentiment['negative_percent']}；"
        f"P/N 值 {sentiment['pn_value']}。",
    ]
    concl += _growth_lines(growth_info)
    if len(keywords) >= 3:
        top3 = "、".join(f"{k['keyword']}（{k['count']}）" for k in keywords[:3])
        concl.append(f"熱門關鍵字 TOP3：{top3}。")
    if kols:
        concl.append(
            f"最具影響力 KOL：{kols[0]['platform']} {kols[0]['channel']}"
            f"（互動 {kols[0].get('engagement', 0):,}）。")
    fw._para(doc, "關鍵結論：", bold=True, space_after=2)
    fw._bullets(doc, concl)

    # ── §2 聲量趨勢 ──────────────────────────────────
    fw._section(doc, "§2　聲量趨勢分析")
    rows = []
    for it in trend["daily"]:
        mark = "★ 高峰" if it["date"] == trend["peak_date"] else ""
        rows.append((it["date"], f"{it['volume']:,}", mark))
    rows.append(("週累積", f"{trend['total']:,}", ""))
    fw._table(doc, ["日期", "聲量（筆）", "趨勢備註"], rows, widths=[2.5, 2.5, 2.0])
    fw._para(doc, "", space_after=2)
    fw._para(doc, "聲量觀察：", bold=True, space_after=2)
    if trend_narrative:
        fw._bullets(doc, list(trend_narrative))
    else:
        fw._bullets(doc, [
            f"高峰日 {trend['peak_date']} 聲量達 {trend['peak_volume']:,} 筆。",
            f"週平均每日 {avg_vol:,} 筆。",
        ])
    if source_dist and source_dist.get("total"):
        parts = [f"{it['type']} {it['count']}篇（{it['percent']}%）"
                 for it in source_dist["items"] if it["count"]]
        fw._para(doc, "來源結構（依熱門文章來源類型，結構參考值）："
                 + "、".join(parts) + "。", size=10, color=C_GRAY, space_before=4)

    # ── §3 情緒分布 ──────────────────────────────────
    fw._section(doc, "§3　情緒分布分析")
    fw._table(doc, ["情緒", "聲量（筆）", "占比"], [
        ("正面", f"{sentiment['positive']:,}", sentiment['positive_percent']),
        ("中立", f"{sentiment['neutral']:,}", sentiment['neutral_percent']),
        ("負面", f"{sentiment['negative']:,}", sentiment['negative_percent']),
        ("合計", f"{trend['total']:,}", "100.00%"),
    ], widths=[2.0, 2.5, 2.5])
    fw._para(doc, "情緒解讀：", bold=True, space_before=4, space_after=2)
    pos_text = _strip_md(ai.get("正面議題", "")).strip() or (
        f"本期正面聲量佔比 {sentiment['positive_percent']}，"
        "主要來自品牌正向討論與產品好評。")
    fw._para(doc, f"【正面（{sentiment['positive_percent']}）】{pos_text}", size=10)
    fw._para(doc, f"【中立（{sentiment['neutral_percent']}）】中立聲量以資訊轉發、"
             "選購詢問型討論為主，情緒中性。", size=10)
    neg_text = _strip_md(ai.get("負面議題", "")).strip() or (
        f"本期負面聲量佔比 {sentiment['negative_percent']}，"
        "請對照熱門文章確認主要負面來源。")
    fw._para(doc, f"【負面（{sentiment['negative_percent']}）】{neg_text}", size=10)

    pn = _safe_float(sentiment['pn_value'])
    if pn >= 2.0:
        note = "P/N 值高於 2.0，整體輿情正向，品牌口碑良好。"
    elif pn >= 1.5:
        note = "P/N 值介於 1.5–2.0，整體輿情偏正面，品牌基本盤穩固。"
    elif pn >= 1.0:
        note = "P/N 值介於 1.0–1.5，正面略優於負面，仍有提升空間。"
    else:
        note = "P/N 值低於 1.0，負面聲量超越正面，建議加強正向議題操作。"
    fw._para(doc, f"P/N 值 {sentiment['pn_value']} 評估：{note}", space_before=4)

    # ── §4 熱門關鍵字 ────────────────────────────────
    fw._section(doc, "§4　熱門關鍵字分析（TOP 20）")
    fw._table(doc, ["排名", "關鍵字", "聲量（筆）"],
              [(k['rank'], k['keyword'], f"{k['count']:,}") for k in keywords[:20]],
              widths=[1.2, 4.0, 1.8])
    kw_events = match_keywords_to_articles(keywords, articles, top_n=8)
    if kw_events:
        fw._para(doc, "關鍵字對應事件（取互動最高之命中文章）：",
                 bold=True, space_before=6, space_after=2)
        ev_rows = []
        for ke in kw_events:
            art = ke.get("article")
            event = (f"[{art['source']}] {_clean(art.get('title', ''), 40)}"
                     if art else "（本期熱門文章未見直接對應）")
            ev_rows.append((ke["keyword"], f"{ke['match_count']} 篇", event))
        fw._table(doc, ["關鍵字", "命中", "對應熱門文章／事件"], ev_rows,
                  widths=[1.8, 1.0, 4.2])
    _caption(doc, table_captions.get("keywords"))

    # ── §5 熱門文章 ──────────────────────────────────
    fw._section(doc, "§5　熱門文章分析（TOP 10）")
    art_rows = []
    for a in articles[:10]:
        art_rows.append((
            a['rank'], a['source'], classify_article(a, brand),
            f"{a['engagement']:,}", f"{a['likes']:,}",
            _clean(a['title'], 34),
        ))
    fw._table(doc, ["排名", "平台", "分類", "互動", "按讚", "標題"], art_rows,
              widths=[0.7, 1.2, 1.2, 1.0, 1.0, 5.0])
    _caption(doc, table_captions.get("articles"))
    if articles:
        top = articles[0]
        obs = [f"互動數最高：[{top['source']}] {_clean(top['channel'], 14)}"
               f"（互動 {top['engagement']:,}、按讚 {top['likes']:,}）。"]
        top_like = max(articles[:10], key=lambda x: x["likes"])
        if top_like["rank"] != top["rank"]:
            obs.append(f"按讚數最高：[{top_like['source']}] "
                       f"{_clean(top_like['channel'], 14)}（按讚 {top_like['likes']:,}）。")
        fw._para(doc, "熱門文章洞察：", bold=True, space_before=4, space_after=2)
        fw._bullets(doc, obs)

    # ── §6 關鍵領袖 KOL ──────────────────────────────
    fw._section(doc, "§6　網路關鍵領袖（KOL）分析 — TOP 10")
    fw._table(doc, ["排名", "平台", "頻道名稱", "發表", "回文", "按讚", "互動總計"],
              [(k['rank'], k['platform'], _clean(k['channel'], 20),
                f"{k.get('posts', 0):,}", f"{k.get('replies', 0):,}",
                f"{k.get('likes', 0):,}", f"{k.get('engagement', 0):,}")
               for k in kols[:10]],
              widths=[0.7, 1.2, 3.0, 0.9, 0.9, 1.1, 1.3])
    _caption(doc, table_captions.get("kol"))
    if kols:
        obs = [f"最具影響力 KOL：{kols[0]['platform']} {kols[0]['channel']}"
               f"（互動總計 {kols[0].get('engagement', 0):,}、"
               f"按讚 {kols[0].get('likes', 0):,}），為本週聲量主要擴散來源。"]
        plats = {}
        for k in kols[:10]:
            plats[k["platform"]] = plats.get(k["platform"], 0) + 1
        if plats:
            tp = max(plats, key=plats.get)
            obs.append(f"KOL 平台分布：{tp} 頻道最多（{plats[tp]} 個），為本週主要擴散平台。")
        fw._para(doc, "KOL 傳播分析：", bold=True, space_before=4, space_after=2)
        fw._bullets(doc, obs)

    # ── §7 核心議題（供分析師補充）───────────────────
    fw._section(doc, "§7　核心議題深度解讀")
    fw._para(doc, "以下依熱門文章 TOP3 自動框架，請分析師補充事件背景與留言走向。",
             color=C_GRAY, space_after=4)
    for i, a in enumerate(articles[:3], 1):
        fw._subsection(doc, f"【議題 {i}】{_clean(a['title'], 50)}")
        fw._para(doc, f"平台：{a['source']}　頻道：{_clean(a['channel'], 20)}　"
                 f"時間：{a['time']}　互動：{a['engagement']:,}　按讚：{a['likes']:,}",
                 size=10, color=C_GRAY, space_after=2)
        fw._bullets(doc, ["事件背景：（請分析師補充）",
                          "輿論走向（正面／負面）：（請補充）",
                          "風險／機會面：（請補充）"], size=10)

    # ── §8 風險觀察 ──────────────────────────────────
    fw._section(doc, "§8　風險觀察與預警")
    neg = sentiment["negative"]
    if neg > 300:
        lvl = f"【高】負面聲量 {neg:,} 筆，超過 300 筆，建議立即查明主要負面來源。"
    elif neg > 100:
        lvl = f"【中】負面聲量 {neg:,} 筆，超過 100 筆，建議持續監控社群討論。"
    else:
        lvl = f"【低】負面聲量 {neg:,} 筆，維持低水位。"
    fw._para(doc, lvl)
    risk = _strip_md(ai.get("風險觀察", "")).strip()
    if risk:
        fw._bullets(doc, [l.strip() for l in risk.split("\n") if l.strip()])

    # ── §9 AI 輿情洞察 ───────────────────────────────
    fw._section(doc, "§9　AI 輿情洞察")
    for key in ["執行摘要", "重點洞察", "正面議題", "負面議題", "風險觀察", "建議行動"]:
        body = _strip_md(ai.get(key, "")).strip()
        if not body:
            continue
        fw._subsection(doc, key)
        lines = [l.strip() for l in body.split("\n") if l.strip()]
        if len(lines) > 1:
            fw._bullets(doc, lines)
        else:
            fw._para(doc, lines[0])

    # ── §10 綜合建議 ─────────────────────────────────
    fw._section(doc, "§10　綜合建議行動")
    act = _strip_md(ai.get("建議行動", "")).strip()
    if act:
        fw._para(doc, "【AI 建議行動】", bold=True, space_after=2)
        fw._bullets(doc, [l.strip() for l in act.split("\n") if l.strip()])
    fw._para(doc, "【分析師補充建議】（短期／中期，請補充）",
             bold=True, color=C_GRAY, space_before=4)

    # ── §11 研究方法 ─────────────────────────────────
    fw._section(doc, "§11　研究方法與資料說明")
    fw._table(doc, ["項目", "內容"], [
        ("資料來源", "KEYPO 大數據關鍵引擎"),
        ("查詢關鍵字", brand),
        ("觀測期間", f"{period}（共 {days} 天）"),
        ("總處理聲量", f"{trend['total']:,} 筆"),
        ("情緒分析", "KEYPO sentidist"),
        ("關鍵字／文章／KOL", "KEYPO hotkw／hotrank／sprdtrnd"),
        ("AI 洞察", "InsightFlow 可切換後端（Ollama／Claude／OpenAI）"),
    ], widths=[2.2, 4.8])

    # ── §12 競品聲量分析（選填）──────────────────────
    if competitor_data:
        fw._section(doc, "§12　競品聲量分析")
        ranked = sorted(
            [{"name": brand, "total": trend["total"],
              "pn": sentiment["pn_value"],
              "pos": sentiment["positive_percent"],
              "neg": sentiment["negative_percent"], "self": True}]
            + [{"name": c["name"], "total": c["trend"]["total"],
                "pn": c["sentiment"]["pn_value"],
                "pos": c["sentiment"]["positive_percent"],
                "neg": c["sentiment"]["negative_percent"], "self": False}
               for c in competitor_data],
            key=lambda x: x["total"], reverse=True)
        fw._table(doc, ["排名", "品牌", "總聲量（筆）", "P/N值", "正面占比", "負面占比"],
                  [(i, r["name"] + ("（本品牌）" if r["self"] else ""),
                    f"{r['total']:,}", str(r["pn"]), r["pos"], r["neg"])
                   for i, r in enumerate(ranked, 1)],
                  widths=[0.8, 2.6, 1.8, 1.2, 1.5, 1.5])

        # 競品事件摘要（AI 敘事）
        comp_caps = table_captions.get("competitors") or {}
        cap_texts = [comp_caps.get(c["name"]) for c in competitor_data
                     if comp_caps.get(c["name"])]
        for ct in cap_texts:
            fw._para(doc, ct, size=10, color=C_GRAY, space_after=1)
        if cap_texts:
            fw._para(doc, AI_CAPTION_NOTE, size=8.5, color=C_GRAY, space_after=2)

        # 比較觀察
        obs = [f"本期聲量最高為「{ranked[0]['name']}」（{ranked[0]['total']:,}筆）。"]
        if ranked[0]["name"] != brand:
            main_rank = next(i for i, r in enumerate(ranked, 1) if r["name"] == brand)
            main_total = next(r["total"] for r in ranked if r["name"] == brand)
            ratio = round(ranked[0]["total"] / main_total, 2) if main_total else None
            ratio_text = f"，約為本品牌的 {ratio} 倍" if ratio else ""
            obs.append(f"{brand} 本期聲量排名第 {main_rank}{ratio_text}。")
        else:
            obs.append(f"{brand} 本期聲量領先所有列入比較的競品。")
        best_pn = max(ranked, key=lambda x: _safe_float(x["pn"]))
        obs.append(f"本期 P/N 值最高為「{best_pn['name']}」（{best_pn['pn']}）。")
        fw._para(doc, "比較觀察：", bold=True, space_before=4, space_after=2)
        fw._bullets(doc, obs)

        # 各競品熱門文章 TOP3
        for c in competitor_data:
            if not c.get("articles"):
                continue
            fw._subsection(doc, f"{c['name']} 熱門文章 TOP3")
            fw._bullets(doc, [
                f"{a['rank']}. [{a['source']}] {_clean(a['title'], 40)}"
                f"（互動 {a['engagement']:,}）" for a in c["articles"][:3]
            ], size=10)

    # ── Footer ───────────────────────────────────────
    fw._para(doc, "", space_before=6)
    fw._para(doc, "本報告由 InsightFlow AI 自動產出；數值與事件描述以 KEYPO 原始資料為準，"
             "§7 核心議題與分析師補充建議請人工填寫。", size=9, color=C_GRAY)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path
