"""
framework_generator.py — 「分析架構與分析腳本」Word 產生器。

對照 Dr.Wu 季報之「分析架構與分析腳本｜提報審核版」排版，輸出一份規劃藍圖文件：
    標題頁 → 一、分析目的 → 二、分析架構（議題×維度矩陣）→
    三、資料來源/工具/口徑（KEYPO 對應＋查詢式原則＋雜訊口徑）→
    四、分析腳本（逐議題：問題×作法/KEYPO×呈現頁型）→
    五、章節結構與頁面對應 → 六、產出與交付

此文件為「規劃藍圖」，不需 KEYPO 資料；由分析計畫（議題／維度／競品／查詢式）驅動。
多數樣板段落（KEYPO 功能對應表、查詢式原則、維度）已內建預設，只需填議題相關內容。

用法：
    from framework_generator import generate_framework
    generate_framework(my_config, "output/架構與腳本.docx")
或直接執行本檔，會用內建範例 config 產生示範文件。
"""
import os
import re
import json

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ── 色盤 ─────────────────────────────────────────────
C_INK    = RGBColor(0x1A, 0x1A, 0x1A)
C_GRAY   = RGBColor(0x5B, 0x64, 0x70)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
SLATE_HEX = "526783"   # 表頭底色
LIGHT_HEX = "F2F4F7"   # 斑馬紋淺色
FONT = "Microsoft JhengHei"


# ── 低階排版工具 ─────────────────────────────────────
def _font(run, size=10.5, bold=False, color=C_INK, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)


def _para(doc, text="", size=10.5, bold=False, color=C_INK,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if text:
        _font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def _red_rule(para, color="C00000", sz=14):
    """在段落下方加一條紅色分隔線（仿母片『黑標題＋紅底線』）"""
    pPr = para._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _section(doc, text):
    """章節大標：黑字粗體 + 紅色底線"""
    p = _para(doc, text, size=15, bold=True, color=C_INK,
              space_before=12, space_after=4)
    _red_rule(p)
    return p


def _subsection(doc, text):
    return _para(doc, text, size=12, bold=True, color=RGBColor(0x2A, 0x3A, 0x4A),
                 space_before=8, space_after=3)


def _bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(str(it)), size=size)


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell(cell, text, size=10, bold=False, color=C_INK,
          bg=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    if bg:
        _shade(cell, bg)
    cell.vertical_alignment = 1  # center
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.space_before = Pt(1)
    _font(para.add_run(str(text)), size=size, bold=bold, color=color)


def _table(doc, headers, rows, widths=None, first_col_header=False):
    """建立表格：首列為標頭（深藍底白字），資料列斑馬紋。
    first_col_header=True 時，每列第一欄也套標頭底色（供矩陣列標題用）。"""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    for j, h in enumerate(headers):
        _cell(tbl.rows[0].cells[j], h, size=10, bold=True, color=C_WHITE,
              bg=SLATE_HEX, align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, row in enumerate(rows):
        cells = tbl.add_row().cells
        zebra = LIGHT_HEX if i % 2 == 1 else None
        for j, val in enumerate(row):
            if first_col_header and j == 0:
                _cell(cells[j], val, size=10, bold=True, color=C_WHITE,
                      bg=SLATE_HEX)
            else:
                _cell(cells[j], val, size=10, bg=zebra)

    if widths:
        for row in tbl.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    return tbl


# ── 內建預設（樣板段落，可被 config 覆寫）─────────────
_DEFAULT_DIMENSIONS = [
    "市場趨勢", "消費者洞察", "品牌 vs 競品", "成長觀察", "可執行建議",
]

_DEFAULT_KEYPO_MAPPING = [
    ("聲量規模與趨勢、QoQ/YoY、月趨勢線", "freqdist", "總聲量筆數＋逐日聲量"),
    ("正/中/負占比、好感度 P/N", "sentidist", "情緒占比"),
    ("新聞/社群/論壇/部落格來源結構", "datadist", "四大來源占比"),
    ("熱門頻道、通路操作", "hotchnl", "頻道排行（已清洗）"),
    ("關鍵領袖 KOL、競品合作", "opleader", "領袖排行（按讚）"),
    ("討論主題、消費者關注點", "hotkwV5", "熱門關鍵字語意雲"),
    ("各頁代表文、網友討論", "textlist", "文章列表（依回文）"),
]

_DEFAULT_QUERY_PRINCIPLES = [
    "共現鎖定：以「主體 & 語境詞」鎖定同篇同時提及者，濾掉離題文。",
    "排除運算：一律用「!」不用「-」；沿用客戶正式「排抽字串」排除抽獎／活動文。",
    "避免裸詞：成分／品類裸詞會灌入電商洗版與跨域雜訊，一律加語境並排除他義。",
]

_DEFAULT_NOISE_NOTES = [
    "本節最關鍵的資料清洗為電商／代購／直播頻道洗版；定稿口徑＝預設代購/直播"
    "頻道排除清單＋客戶正式排抽字串，只保留真實消費者、KOL、媒體與品牌官方討論。",
    "YoY 判讀：若同期含品牌活動高峰或單串活動留言灌水造成失真，於頁面上如實拆解，"
    "輔以 QoQ 與近月趨勢判讀。",
    "殘留誠實標註：成分類主題絕對聲量仍含少量商業貼文，以趨勢形狀＋乾淨代表文判讀；"
    "附錄附查詢式與口徑供覆核。",
]

_SCRIPT_HEADERS = ["要回答的問題", "作法／KEYPO", "呈現頁型"]


# ══════════════════════════════════════════════════════
# 主函式
# ══════════════════════════════════════════════════════
def generate_framework(config, out_path):
    cfg = dict(config)
    dimensions = cfg.get("dimensions", _DEFAULT_DIMENSIONS)

    doc = Document()
    # 全文預設字型
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # ── 標題區 ────────────────────────────────────────
    t = _para(doc, cfg.get("report_title", "網路輿情分析"), size=22, bold=True,
              color=C_INK, space_after=2)
    _red_rule(t, sz=20)
    _para(doc, cfg.get("subtitle", "分析架構與分析腳本｜提報審核版"),
          size=13, color=C_RED, space_after=8)

    meta = [
        ("觀測期間", cfg.get("period", "")),
        ("對照基準", cfg.get("compare", "")),
        ("資料來源", cfg.get("data_source", "")),
        ("指定議題", "｜".join(cfg.get("topics", []))),
        ("製作單位", cfg.get("producer", "")),
        ("文件版本", cfg.get("doc_version", "")),
    ]
    for label, val in meta:
        if not val:
            continue
        p = _para(doc, "", space_after=2)
        _font(p.add_run(f"{label}："), size=10, bold=True, color=C_GRAY)
        _font(p.add_run(str(val)), size=10, color=C_INK)

    # ── 一、分析目的 ─────────────────────────────────
    _section(doc, "一、分析目的與本季調整重點")
    if cfg.get("objective"):
        _para(doc, cfg["objective"], size=10.5)
    if cfg.get("core_questions"):
        _para(doc, "三大議題共同要回答的核心問題：", size=10.5, bold=True,
              space_after=2)
        _bullets(doc, cfg["core_questions"])
    if cfg.get("season_note"):
        _para(doc, cfg["season_note"], size=10.5, space_before=4)

    # ── 二、分析架構（議題 × 維度矩陣）────────────────
    _section(doc, "二、分析架構：議題 × 分析維度")
    _para(doc, "以「議題（列）× 分析維度（欄）」矩陣統一結構，"
               "確保各議題深度一致、可橫向比較：", size=10.5)
    matrix = cfg.get("matrix", {})
    headers = ["議題＼維度"] + dimensions
    rows = []
    for topic in cfg.get("topics", []):
        cell_map = matrix.get(topic, {})
        rows.append([topic] + [cell_map.get(d, "—") for d in dimensions])
    if rows:
        _table(doc, headers, rows, first_col_header=True)

    # ── 三、資料來源、工具與觀測口徑 ──────────────────
    _section(doc, "三、資料來源、工具與觀測口徑")
    _subsection(doc, "3-1　KEYPO 功能對應（每個分析問題用哪支 API）")
    _table(doc, ["分析用途", "KEYPO 功能", "輸出"],
           cfg.get("keypo_mapping", _DEFAULT_KEYPO_MAPPING),
           widths=[3.6, 1.5, 2.1])
    _subsection(doc, "3-2　查詢式設計原則")
    _bullets(doc, cfg.get("query_principles", _DEFAULT_QUERY_PRINCIPLES))
    _subsection(doc, "3-3　雜訊處理與觀測口徑（誠實說明，供覆核）")
    _bullets(doc, cfg.get("noise_notes", _DEFAULT_NOISE_NOTES))

    # ── 四、分析腳本（逐議題）─────────────────────────
    _section(doc, "四、分析腳本（逐議題）")
    for script in cfg.get("scripts", []):
        _subsection(doc, script.get("title", ""))
        rows = [(r[0], r[1], r[2]) for r in script.get("rows", [])]
        if rows:
            _table(doc, _SCRIPT_HEADERS, rows, widths=[2.6, 3.1, 1.5])

    # ── 五、報告章節結構與頁面對應 ────────────────────
    if cfg.get("chapters"):
        _section(doc, "五、報告章節結構與定稿頁面對應")
        if cfg.get("visual_note"):
            _para(doc, cfg["visual_note"], size=10.5)
        _table(doc, ["章節", "內容", "頁數（投影片序）"],
               cfg["chapters"], widths=[1.9, 4.3, 1.5])

    # ── 六、產出與交付 ───────────────────────────────
    if cfg.get("deliverables"):
        _section(doc, "六、產出與交付")
        _bullets(doc, cfg["deliverables"])

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


# ══════════════════════════════════════════════════════
# AI 依品牌自動產生 config（內容 AI 擬、架構固定套用本範本）
# ══════════════════════════════════════════════════════
import ai_backend

_DEFAULT_CORE_QUESTIONS = [
    "市場在成長還是萎縮？季節性與關鍵節點為何？（趨勢）",
    "消費者在討論什麼、為什麼買、卡在哪？正負評與使用情境／痛點為何？（洞察）",
    "本品牌與主要競品的聲量位置、操作手法與內容差異為何？（競品）",
    "哪些品牌／成分／產品聲量成長最快、原因是什麼？（成長）",
    "品牌下一步可以怎麼做？（可執行建議）",
]

_DEFAULT_DELIVERABLES = [
    "各議題市場動態觀察（Word）── 同步提交。",
    "分析架構與分析腳本（Word）── 即本文件。",
    "實作過程說明（TXT）── 記錄取數規則、口徑、清洗與 PPT 技術規範。",
    "定稿交付：季報 PPTX（含 PDF）＋文章列表 xlsx＋關鍵字查詢式清單 xlsx。",
]


def _extract_json(text):
    """從模型回應抽出 JSON（去 markdown 圍欄、取首個 { 到末個 }）"""
    text = str(text).strip()
    text = re.sub(r'^```(?:json)?', '', text).strip()
    text = re.sub(r'```$', '', text).strip()
    a, b = text.find('{'), text.rfind('}')
    if a < 0 or b <= a:
        raise ValueError("AI 未回傳可解析的 JSON")
    snippet = text[a:b + 1]
    snippet = re.sub(r',(\s*[}\]])', r'\1', snippet)   # 去除尾逗號（本地模型常見）
    return json.loads(snippet)


def _build_chapters(topics, brand):
    """依議題數自動組出章節↔頁碼對應（deterministic）"""
    plan = [
        ("封面／目錄", "標題＋觀測期間；章節目錄", 2),
        ("01 研究方法", "KEYPO 引擎調查方法＋名詞說明＋資料口徑（排除抽獎與洗版頻道）", 5),
        ("本季重點摘要", f"各議題定調＋{brand} 三指標卡（排抽聲量／SOV／好感度）", 1),
    ]
    for i, t in enumerate(topics, start=2):
        name = t.split("、", 1)[-1] if "、" in t else t
        plan.append((f"{i:02d} {name}",
                     "趨勢→洞察關鍵字→好感度正/負→競品聲量→小結與建議", 8))
    n = len(topics)
    plan.append((f"{n + 2:02d} 品牌競爭格局",
                 "含抽排行→排抽排行→社群操作→熱門產品事件→KOL→排抽總表→好感度定位", 8))
    plan.append((f"{n + 3:02d} 總結與建議", "總結（品牌共同課題）＋各議題 Insights×建議 總表", 3))
    plan.append((f"{n + 4:02d} 附錄", "關鍵字查詢式清單＋已刪除貼文補充；末頁 Thank You", 4))

    rows, start = [], 1
    for name, content, cnt in plan:
        end = start + cnt - 1
        rng = f"{cnt} 頁（{start}–{end}）" if cnt > 1 else f"1 頁（{start}）"
        rows.append((name, content, rng))
        start = end + 1
    return rows


def build_config_with_ai(brand, industry="", competitors=None, n_topics=3,
                         period="", producer="", timeout=180):
    """用 AI 後端（AI_BACKEND：ollama/claude/openai）依品牌自動擬出議題、
    議題×維度矩陣、逐議題腳本與競品；架構與樣板段落固定套用本範本。
    回傳可直接餵給 generate_framework() 的 config。"""
    dims = _DEFAULT_DIMENSIONS
    comp_hint = ("、".join(competitors) if competitors
                 else "（請依品牌類別自行推斷主要競品）")

    prompt = f"""你是一位資深品牌輿情分析顧問，正在為「{brand}」規劃一份網路輿情季報的分析架構。
品牌類別／背景：{industry or "（請依品牌名稱推斷所屬產業與消費者關注點）"}
主要競品（可參考或補充）：{comp_hint}

請規劃 {n_topics} 個最值得分析的「市場議題」——不是單一品牌，而是該品牌所處市場中，消費者正在熱議且與品牌高度相關的題目。

【矩陣（matrix）】為每個議題的「五個維度」各填「一句話：這個議題在這個維度要看什麼」。
維度固定為：{"、".join(dims)}（矩陣的鍵必須與這五個完全一致，五格都要填、不可留空）。
矩陣格「只寫要觀察的內容」，不要寫 KEYPO 端點或頁型。範例（油類保養議題）：
  市場趨勢→「以油養膚/精華油月趨勢、季節性」
  消費者洞察→「正負評、使用情境（夜間修護/換季）、痛點（致痘/黏膩/味道）」
  品牌 vs 競品→「本品牌精華油 vs The Ordinary/Kiehl's/克蘭詩」
  成長觀察→「近期成長較快品牌與原因」
  可執行建議→「溝通與內容方向」

【腳本（scripts）】為每個議題列 3～5 個「要回答的問題」，每題對應「作法／KEYPO」與「呈現頁型」。
KEYPO 端點（只在腳本的『作法／KEYPO』欄出現）：freqdist（趨勢）、sentidist（好感度）、datadist（來源結構）、hotchnl（頻道）、opleader（KOL）、hotkwV5（關鍵字語意雲）、textlist（代表文）。
呈現頁型例如：趨勢折線頁、洞察關鍵字頁、好感度正/負熱議雙頁、競品對比頁、成長頁、單品深挖頁。

只輸出 JSON，不要任何說明文字、不要 markdown 圍欄，格式：
{{
  "topics": ["一、議題名", "二、議題名"],
  "matrix": {{"一、議題名": {{"{dims[0]}":"該維度要看什麼（一句話）", "{dims[1]}":"...", "{dims[2]}":"...", "{dims[3]}":"...", "{dims[4]}":"..."}}}},
  "scripts": [{{"title":"4-1　議題名", "rows": [["要回答的問題","freqdist 月趨勢＋季節性","趨勢折線頁"]]}}],
  "competitors": ["競品1","競品2"]
}}
務必使用繁體中文（台灣用語），內容要具體貼合「{brand}」與其市場，不要空泛。"""

    raw = ai_backend.generate(prompt, timeout=timeout, max_tokens=4000)
    data = _extract_json(raw)
    topics = data["topics"]

    config = {
        "report_title": f"{brand}｜網路輿情 市場議題分析規劃",
        "subtitle": "分析架構與分析腳本（規劃藍圖｜AI 初擬，待分析師核定）",
        "period": period,
        "compare": "對照：上一季 QoQ、去年同期 YoY；趨勢線近 12 個月",
        "data_source": "KEYPO 大數據關鍵引擎（新聞／社群／論壇／部落格全網）＋ Google／Web 佐證",
        "producer": producer,
        "doc_version": f"分析架構 v1.0（AI 初擬｜後端 {ai_backend.backend_name()}）",
        "topics": topics,
        "objective": (
            "【文件定位】本文件為前瞻性的「市場議題分析規劃」——規劃下一階段值得深入的"
            "市場議題與分析腳本，供分析師核定；與本期單一品牌的輿情數據報告（簡報 PPT）"
            "用途不同：本文件回答「接下來要分析哪些市場議題」，簡報則呈現「該品牌本期"
            "實際的聲量與輿情數據」。　"
            f"延續 {brand} 網路輿情追蹤，本階段聚焦下列 {len(topics)} 大市場議題，"
            "並在市場數據與聲量分析之外加重 Insights 與 Actionable Recommendations；"
            "每一議題以相同分析骨架處理，段末固定產出「洞察」與「可執行建議」。"
        ),
        "core_questions": _DEFAULT_CORE_QUESTIONS,
        "matrix": data["matrix"],
        "scripts": data["scripts"],
        "chapters": _build_chapters(topics, brand),
        "deliverables": _DEFAULT_DELIVERABLES,
        "visual_note": (
            "定稿視覺＝品牌簡報母片（自動頁尾 logo＋版權＋頁碼；內容頁黑標題＋短紅底線＋"
            "右上觀測期間；圖表 KEYPO 配色；全檔微軟正黑體）。"
        ),
    }
    if data.get("competitors"):
        config["season_note"] = (
            "本季新增「品牌競爭格局」章：回到品牌層級比較 "
            f"{brand} 與主要競品（{'、'.join(data['competitors'])}）的聲量排行"
            "（含抽獎／排除抽獎兩口徑）、市佔（SOV）、社群操作、熱門產品事件、"
            "關鍵領袖與好感度定位。"
        )
    return config


# ══════════════════════════════════════════════════════
# 內建示範 config（對照 Dr.Wu 範本；實際使用時替換成你的計畫）
# ══════════════════════════════════════════════════════
SAMPLE_CONFIG = {
    "report_title": "○○品牌 網路輿情季報 2026 第二季",
    "subtitle": "分析架構與分析腳本｜提報審核版",
    "period": "2026/04/01 – 06/30",
    "compare": "對照：2026 Q1 之 QoQ、2025 Q2 之 YoY；趨勢線 2025/07–2026/06 近 12 個月",
    "data_source": "KEYPO 大數據關鍵引擎（新聞／社群／論壇／部落格全網）＋ Google／Web 佐證",
    "producer": "○○大數據股份有限公司",
    "doc_version": "分析架構 v1.0",
    "topics": ["一、油類保養", "二、膠原飲粉", "三、抗老/PDRN/外泌體"],
    "objective": (
        "延續品牌季度網路輿情追蹤，本季依客戶需求聚焦上述三大議題，並在市場數據與"
        "聲量分析之外加重 Insights 與 Actionable Recommendations。每一議題均以相同"
        "分析骨架處理，段末固定產出「洞察」與「可執行建議」，使報告可直接對接品牌"
        "行銷、產品溝通與內容策略。"
    ),
    "core_questions": [
        "市場在成長還是萎縮？季節性與關鍵節點為何？（趨勢）",
        "消費者在討論什麼、為什麼買、卡在哪？正負評與使用情境／痛點為何？（洞察）",
        "本品牌與主要競品的聲量位置、操作手法與內容差異為何？（競品）",
        "哪些品牌／成分聲量成長最快、原因是什麼？（成長）",
        "品牌下一步可以怎麼做？（可執行建議）",
    ],
    "season_note": (
        "本季新增（依客戶指定競品名單）：於三大議題之後加入「品牌競爭格局」章，"
        "回到品牌層級比較本品牌與主要競品的聲量排行（含抽獎／排除抽獎兩口徑）、"
        "市佔（SOV）、社群操作、熱門產品事件、關鍵領袖與好感度定位。"
    ),
    "matrix": {
        "一、油類保養": {
            "市場趨勢": "以油養膚/精華油月趨勢、季節性",
            "消費者洞察": "正負評、使用情境（夜間修護/換季）、痛點（致痘/黏膩/味道）",
            "品牌 vs 競品": "本品牌精華油 vs The Ordinary / Kiehl's / 克蘭詩",
            "成長觀察": "成長品牌與原因",
            "可執行建議": "溝通／內容方向",
        },
        "二、膠原飲粉": {
            "市場趨勢": "膠原保健、膠原胜肽飲、上錠下飲趨勢",
            "消費者洞察": "最有感溝通方式、代言/科技訴求、好感度",
            "品牌 vs 競品": "本品牌胜肽飲 vs 蜜露珂娜 / M2 / TKLAB",
            "成長觀察": "成長品牌與原因",
            "可執行建議": "溝通／內容方向",
        },
        "三、抗老/PDRN/外泌體": {
            "市場趨勢": "抗老精華、PDRN、外泌體成分趨勢",
            "消費者洞察": "TA 是否年輕化、關注點、使用情境",
            "品牌 vs 競品": "本品牌抗老精華 vs 理膚AB / 赫蓮娜 / 未來美PDRN",
            "成長觀察": "成分接棒與爆紅原因",
            "可執行建議": "溝通／內容方向",
        },
    },
    "scripts": [
        {
            "title": "4-1　油類保養市場",
            "rows": [
                ("整體聲量與討論趨勢？",
                 "freqdist 月趨勢＋季節性（油類為秋冬題，觀察 Q2 走勢）", "趨勢折線頁"),
                ("消費者正負面與使用情境／痛點？",
                 "sentidist＋hotkwV5＋textlist（致痘/黏膩 vs 修護/清爽）",
                 "洞察關鍵字頁＋好感度正/負熱議雙頁"),
                ("本品牌與競品操作與內容差異？",
                 "本品牌精華油 vs The Ordinary/Kiehl's/克蘭詩 freqdist＋opleader",
                 "對比頁"),
                ("近期成長較快品牌及原因？",
                 "產品 QoQ＋Google 佐證（檔期/代言/新品）", "成長頁"),
            ],
        },
        {
            "title": "4-2　台灣膠原市場（膠原飲／膠原粉）",
            "rows": [
                ("市場趨勢與討論概況？",
                 "膠原保健 freqdist＋抽獎占比＋來源結構", "趨勢＋結構頁"),
                ("本品牌胜肽飲上市定位？",
                 "本品牌膠原產品月趨勢（定案查詢式）＋品項聲量排序", "單品深挖頁"),
                ("消費者最有感的溝通方式與內容？",
                 "hotkwV5＋textlist（代言/科技/分子量/風味/劑量）", "洞察頁"),
                ("競品操作與成長？",
                 "蜜露珂娜/M2/TKLAB freqdist＋opleader＋Google", "競品＋成長頁"),
            ],
        },
    ],
    "chapters": [
        ("封面／目錄", "標題＋觀測期間；章節目錄", "2 頁（1–2）"),
        ("01 研究方法", "KEYPO 引擎調查方法＋名詞說明＋資料口徑（排除抽獎與洗版頻道）",
         "5 頁（3–7）"),
        ("本季重點摘要", "三議題定調＋品牌三指標卡（排抽聲量／SOV／好感度）", "1 頁（8）"),
        ("02 油類保養", "趨勢→各渠道聲量→洞察關鍵字→好感度正/負→競品→小結與建議",
         "8 頁（9–16）"),
        ("03 膠原飲粉", "趨勢→洞察關鍵字→本品牌胜肽飲趨勢→競品→小結與建議",
         "7 頁（17–23）"),
        ("04 抗老/PDRN/外泌體",
         "趨勢→洞察→PDRN vs 外泌體→好感度正/負→競品→小結與建議", "10 頁（24–33）"),
        ("05 品牌競爭格局",
         "含抽排行→排抽排行→社群操作→熱門產品事件→KOL→排抽總表→好感度定位",
         "8 頁（34–41）"),
        ("06 總結與建議", "總結（品牌共同課題）＋四議題 Insights×建議 總表", "3 頁（42–44）"),
        ("07 附錄", "關鍵字查詢式清單＋已刪除貼文補充；末頁 Thank You", "4 頁（45–48）"),
    ],
    "visual_note": (
        "定稿視覺＝品牌簡報母片（自動頁尾：logo＋版權＋頁碼；內容頁黑標題＋短紅底線＋"
        "右上觀測期間；圖表 KEYPO 配色；全檔微軟正黑體）。"
    ),
    "deliverables": [
        "三大議題市場動態觀察（Word）── 同步提交。",
        "分析架構與分析腳本（Word）── 即本文件。",
        "實作過程說明（TXT）── 記錄取數規則、口徑、清洗與 PPT 技術規範。",
        "定稿交付：季報 PPTX（含 PDF）＋文章列表 xlsx＋關鍵字查詢式清單 xlsx。",
    ],
}


if __name__ == "__main__":
    print("=" * 50)
    print("分析架構與分析腳本 產生器")
    print(f"AI 後端：{ai_backend.backend_name()}"
          "（可用環境變數 AI_BACKEND 切換 ollama / claude / openai）")
    print("=" * 50)
    brand = input("\n請輸入品牌名稱（直接 Enter 使用內建示範）：").strip()

    if not brand:
        out = generate_framework(
            SAMPLE_CONFIG, "output/_framework_sample/分析架構與分析腳本_示範.docx")
        print(f"\n已產出示範文件：{out}")
    else:
        industry = input("品牌產業／背景（選填，有助 AI 精準推斷）：").strip()
        period = input("觀測期間（選填，例 2026/04/01 – 06/30）：").strip()
        print("\nAI 依品牌自動擬定議題、矩陣、腳本、競品中…")
        config = build_config_with_ai(brand, industry=industry, period=period)
        out = generate_framework(
            config, f"output/{brand}/市場議題分析規劃_{brand}.docx")
        print(f"\n✅ 已產出：{out}")
        print("（內容為 AI 初擬，請分析師核定後定稿）")
